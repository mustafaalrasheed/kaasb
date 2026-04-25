"""Minimal markdown → .docx converter for the engineering handoff.

Handles: ATX headings, paragraphs, ordered/unordered lists (one level),
fenced code blocks, pipe tables, bold (**), italics (*), inline code (`),
links [text](url), blockquotes (>), horizontal rules (---).
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor


INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
ITALIC = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_inline(paragraph, text: str) -> None:
    """Render inline markdown (bold/italic/code/link) into a paragraph."""
    pos = 0
    tokens: list[tuple[int, int, str, str, str]] = []  # (start, end, kind, text, href)

    for m in LINK.finditer(text):
        tokens.append((m.start(), m.end(), "link", m.group(1), m.group(2)))
    for m in BOLD.finditer(text):
        tokens.append((m.start(), m.end(), "bold", m.group(1), ""))
    for m in INLINE_CODE.finditer(text):
        tokens.append((m.start(), m.end(), "code", m.group(1), ""))
    for m in ITALIC.finditer(text):
        tokens.append((m.start(), m.end(), "italic", m.group(1), ""))

    tokens.sort(key=lambda t: (t[0], -t[1]))

    cleaned: list[tuple[int, int, str, str, str]] = []
    last_end = -1
    for tok in tokens:
        if tok[0] >= last_end:
            cleaned.append(tok)
            last_end = tok[1]

    for start, end, kind, content, href in cleaned:
        if start > pos:
            paragraph.add_run(text[pos:start])
        run = paragraph.add_run(content)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif kind == "link":
            run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
            run.underline = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, cell_text.strip())
            if r == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def parse_table_block(lines: list[str], i: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while i < len(lines) and "|" in lines[i] and lines[i].strip():
        line = lines[i].strip().strip("|")
        cells = [c.strip() for c in line.split("|")]
        if all(re.match(r"^:?-+:?$", c) for c in cells if c):
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if "CodeBlock" not in [s.name for s in styles]:
        cb = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        cb.font.name = "Consolas"
        cb.font.size = Pt(9)
        cb.paragraph_format.left_indent = Pt(18)
        cb.paragraph_format.space_before = Pt(2)
        cb.paragraph_format.space_after = Pt(2)

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        if line.lstrip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                for code_line in code_buf:
                    p = doc.add_paragraph(style="CodeBlock")
                    p.add_run(code_line)
                in_code = False
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            heading_text = re.sub(r"^[^\w؀-ۿ]+", "", heading_text)
            h = doc.add_heading(level=min(level, 4))
            add_inline(h, heading_text)
            i += 1
            continue

        # Tables
        if "|" in line and i + 1 < len(lines) and re.search(r"\|\s*:?-", lines[i + 1]):
            rows, i = parse_table_block(lines, i)
            render_table(doc, rows)
            continue

        # Blockquote
        if stripped.startswith(">"):
            quoted = stripped.lstrip("> ").strip()
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, quoted)
            i += 1
            continue

        # Ordered list
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            content = m.group(3)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, content)
            i += 1
            continue

        # Unordered list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            content = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content)
            i += 1
            continue

        # Paragraph (collect continuation lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip():
                break
            if re.match(r"^(#{1,6}\s|>|\d+\.\s|\s*[-*]\s|```)", nxt):
                break
            if "|" in nxt and i + 1 < len(lines) and re.search(r"\|\s*:?-", lines[i + 1] if i + 1 < len(lines) else ""):
                break
            para_lines.append(nxt.strip())
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(para_lines))

    doc.save(str(docx_path))


if __name__ == "__main__":
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    convert(src, dst)
    print(f"Wrote {dst} ({dst.stat().st_size:,} bytes)")
